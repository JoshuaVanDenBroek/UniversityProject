from docx import Document
from docx.shared import Pt
import easygui

def encrypt():  #encryption
    Plaintext = ''
    Keyword = (InputtedKeyword)
    
    Plaintext = (plain_input_para).strip() #User plaintext input
    print(plain_input_para)
    Temp = Plaintext

    KeywordList = list(Keyword) #Keyword converted into list of letters

    KeywordDinaryForm = list()
    KeywordDinaryForm1 = list()
    Counter = 0
    
    for x in KeywordList:
        KeywordDinaryForm.append(ord(x)) #Keyword converted into denary form (ascii) (added to new list)

    print(KeywordDinaryForm)
    for x in KeywordDinaryForm:
        if (x+Counter) > 126:
            howlarge = (x+Counter)-126
            newData = (32+howlarge)
            KeywordDinaryForm1.append(newData) #Keyword denary value as counter added each repeat
            
        else:
            KeywordDinaryForm1.append(x+Counter) #Keyword denary value as counter added each repeat
            
        Counter += 1
    KeywordList.clear() #Original keyword list is cleared

    print(KeywordDinaryForm1)
    
    for x in KeywordDinaryForm1:
        KeywordList.append(chr(x)) #Keyword converted back into characters from denary

    Keyword = ""

    for x in KeywordList:
        Keyword = Keyword+x
    
    LetterNum=0
        
    for x in Plaintext:
        LetterNum = LetterNum+1

    if(LetterNum % 2) != 0:
        Plaintext = Plaintext+'X' #If the new plaintext is odd an 'X' is added to the end to make it even
    
    Text = list(Plaintext) #plaintext is converted to a list

    LetterNum=0
    for x in Text:
        LetterNum = LetterNum+1

    Target = (len(Text)/2)
    Cipher1 = list()
    Cipher2 = list()
        
    Right = (len(Text)-1)
    Left = 0
    z = 0
    while z < Target: #This while loop is used to scramble up the plaintext
        Cipher1.append(Text[Left]) #one letter is taken from the left of the plaintext and added to a new list
        Left = Left+1
        Cipher1.append(Text[Right]) #one letter is taken from the right of the plaintext and added to a new list
        Right = Right-1
        z += 1 #This is then repeated until z equals target which is half of the length of list text

    Right = (len(Text)-1)
    Left = 0
    z = 0
    while z < Target: #The scrambling of the plaintext is repeated twice to increase the scramble
        Cipher2.append(Cipher1[Left])
        Left = Left+1
        Cipher2.append(Cipher1[Right])
        Right = Right-1
        z += 1

    Cipher3 = list()
    Cipher4 = list()
    Repeats = 0

    for x in range(len(Keyword)):
        Repeats = Repeats + (ord(Keyword[x])-64)

    while Repeats != 0:
        Repeats = Repeats-1
        Cipher3.clear()

        Amount = (len(Cipher2))
        y = 0

        for x in range(len(Cipher2)): #Next we take the scrambled plaintext and encrypt it using the keyword
            if x != Amount:
                if y != len(Keyword)-1: #Here we are going through the letters within the keyword over and over
                    y += 1
                else:
                    y = 0
                Value = (ord(Keyword[y]))
                Temp = (ord(Cipher2[x])+Value)

                if(Temp > 126):
                    while(Temp > 126):
                        Temp = ((Temp-126)+32)
                    Cipher3.append(Temp)
                else:
                    Cipher3.append(ord(Cipher2[x])+Value)   #Each of the characters of the keyword have a denary value
                                                        #and that is added to the value of the current letter we are encrypting in the plaintext list
        Cipher2.clear()
        for x in range(len(Cipher3)):
            Cipher2.append(chr(Cipher3[x]))

    for x in range(len(Cipher3)):
        Cipher4.append(chr(Cipher3[x])) #Once encrypted the denary list now has to be converted back to characters again (ascii) and then added to a new list

    Ciphertext = ""

    for x in Cipher4:
        Ciphertext = Ciphertext + x

    Input = {plain_input_para: Ciphertext}

    for i in Input:
        for p in doc.paragraphs:
            for r in paragraph.runs:
                if r.text.find(i)>=0 and r.font.color.rgb == (255,0,0):
                    r.text=r.text.replace(i,Input[i])
                    r.font.color.rgb == (255,0,0)

#----------------------------------------------------------------------#
Sure = False
while Sure == False:
    file = easygui.fileopenbox()
    if(file[-4:] == 'docx'):
        Sure = (easygui.ynbox('Are you sure you want to encrypt this file?', 'Are you sure?', ('Yes', 'No')))
    else:
        easygui.msgbox('Please select a Microsoft Word Document Only (.docx)', 'Error')

InputtedKeyword = ''

while (InputtedKeyword.isalpha() == False):
    InputtedKeyword = (easygui.passwordbox('Please input your keyword (letters ONLY): '))
InputtedKeyword=(InputtedKeyword.upper()).strip()

doc=Document(file)

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.color.rgb == (255,0,0):
            print(run.text)
            plain_input_para = (run.text)
            encrypt()
            doc.save(file[:-5]+(" (Encrypted).docx"))
            
